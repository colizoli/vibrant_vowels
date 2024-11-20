#!/usr/bin/python2
# -*- coding: utf-8 -*-
"""
Vibrant Vowels: Vowel-grapheme to color code for Word documents (DOCX)
O.Colizoli 2024
Python 3.9
"""
import os, time, itertools, copy
import docx
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
import pandas as pd
import numpy as np

# from IPython import embed as shell # for Olympia's debugging only

'''
Notes
-------------
This script will loop through individual letters and reformat each occurrence of a letter with a specific color (RGB);
If 'change_font' is True, the script will replace the entire document's font name and font size. 
The reformatted document is saved as a new document. 
Fonts need to be already installed on the computer.
---
Books are found in the directory: os.path.join(os.getcwd(), 'books')
The name of book input at the prompt should exclude the 'docx' file extension: "books/{}.docx".format(book_name)
The new colored version will be saved as: "books/{}_vibrant_vowels.docx".format(book_name)
---
Letters and colors to be changed are imported from a CSV file in the directory: os.path.join(os.getcwd(), 'colors')
The CSV file with colors needs to have the following columns: 'letter', 'r', 'g', 'b'
Letters are case-sensitive.
---
The following packages need to be installed: python-docx, pandas, numpy
>> pip install python-docx pandas numpy
---
To run:
>> python color_text_vibrant_vowels.py
You will be prompted with the book name and whether you want to change the font type and size.
---
Character formatting is applied at the docx.text.run.Run level. 
The script can be adjusted to change the font typeface, size, bold, italic, 
and underline of single letters or the whole document.
A Run object has a read-only font property providing access to a Font object. 
A run's Font object (docx.text.run.Run.font) provides properties for getting and setting the character formatting for that run.
E.g. current_run.font.color.rgb = RGBColor(r, g, b) 
---
The function for isolating individual letters as runs, isolate_run(), was taken from here:
See: https://github.com/python-openxml/python-docx/issues/980
'''

def isolate_run(paragraph, start, end):
    """Return docx.text.run object containing only `paragraph.text[start:end]`.
    
    Notes
    -----
    Runs are split as required to produce a new run at the `start` that ends at `end`.
    Runs are unchanged if the indicated range of text already occupies its own run. The
    resulting run object is returned.

    `start` and `end` are as in Python slice notation. For example, the first three
    characters of the paragraph have (start, end) of (0, 3). `end` is not the index of
    the last character. These correspond to `match.start()` and `match.end()` of a regex
    match object and `s[start:end]` of Python slice notation.
    
    https://github.com/python-openxml/python-docx/issues/980
    """
    rs = tuple(paragraph._p.r_lst)

    def advance_to_run_containing_start(start, end):
        """Return (r_idx, start, end) triple indicating start run and adjusted offsets.

        The start run is the run the `start` offset occurs in. The returned `start` and
        `end` values are adjusted to be relative to the start of `r_idx`.
        """
        # --- add 0 at end so `r_ends[-1] == 0` ---
        r_ends = tuple(itertools.accumulate(len(r.text) for r in rs)) + (0,)
        r_idx = 0
        while start >= r_ends[r_idx]:
            r_idx += 1
        skipped_rs_offset = r_ends[r_idx - 1]
        return rs[r_idx], r_idx, start - skipped_rs_offset, end - skipped_rs_offset

    def split_off_prefix(r, start, end):
        """Return adjusted `end` after splitting prefix off into separate run.

        Does nothing if `r` is already the start of the isolated run.
        """
        if start > 0:
            prefix_r = copy.deepcopy(r)
            r.addprevious(prefix_r)
            r.text = r.text[start:]
            prefix_r.text = prefix_r.text[:start]
        return end - start

    def split_off_suffix(r, end):
        """Split `r` at `end` such that suffix is in separate following run."""
        suffix_r = copy.deepcopy(r)
        r.addnext(suffix_r)
        r.text = r.text[:end]
        suffix_r.text = suffix_r.text[end:]

    def lengthen_run(r, r_idx, end):
        """Add prefixes of following runs to `r` until `end` is reached."""
        while len(r.text) < end:
            suffix_len_reqd = end - len(r.text)
            r_idx += 1
            next_r = rs[r_idx]
            if len(next_r.text) <= suffix_len_reqd:
                # --- subsume next run ---
                r.text = r.text + next_r.text
                next_r.getparent().remove(next_r)
                continue
            if len(next_r.text) > suffix_len_reqd:
                # --- take prefix from next run ---
                r.text = r.text + next_r.text[:suffix_len_reqd]
                next_r.text = next_r.text[suffix_len_reqd:]

    r, r_idx, start, end = advance_to_run_containing_start(start, end)
    end = split_off_prefix(r, start, end)

    # --- if run is longer than isolation-range we need to split-off a suffix run ---
    if len(r.text) > end:
        split_off_suffix(r, end)
    # --- if run is shorter than isolation-range we need to lengthen it by taking text
    # --- from subsequent runs
    elif len(r.text) < end:
        lengthen_run(r, r_idx, end)

    return docx.text.run.Run(r, paragraph)
  

''' RUN '''      
if __name__ == "__main__":
    
    book_name = input("Name of book: ")
    change_font = int(input("Change Font? (1 for Yes, 0 for No): "))
    
    if change_font:
        replace_font = input("Font name: ")
        replace_size = np.float32(input("Font size: "))

    in_book_filename = os.path.join('books', "{}.docx".format(book_name)) # original
    out_book_filename = os.path.join('books', "{}_vibrant_vowels.docx".format(book_name)) # save as new
    
    # Define letters and colors to replace
    df = pd.read_csv(os.path.join('colors', 'vibrant_vowels_colors.csv')) # the CSV file with the 'letters' and 'r', 'g', 'b' values

    letters = df['letter'] # vowels including y
    colors_r = df['r'] # red value for RGB code
    colors_g = df['g'] # green value for RGB code
    colors_b = df['b'] # blue value for RGB code
    
    t0 = time.time() # measure run time (not optimized, just curious)

    # First, run the letter-by-letter search and replace loop
    doc = Document(in_book_filename)
    for idx_letter,letter in enumerate(letters): # loop over letters
    
        print('Searching for "{}"...'.format(letter))
        print('Number of paragraphs: {}'.format(len(doc.paragraphs)))
    
        for p_idx,paragraph in enumerate(doc.paragraphs):
            # print('Paragraph {}'.format(p_idx))
        
            for start in range(len(paragraph.text)): # isolate runs that are 1 character in length only
                end = start + 1 # the 1 indicates the step size to search for strings, i.e., only 1 unit from start will equal a single character
                current_run = isolate_run(paragraph, start, end)
            
                if current_run.text == letter: # only change the color of letters in the CSV file (case-sensitive)
                    current_run.font.color.rgb = RGBColor(int(colors_r[idx_letter]), int(colors_g[idx_letter]), int(colors_b[idx_letter]))   
                    # If you want to change other formatting options of the individual letters, you can specify that here:
                    # current_run.font.size
                    # current_run.font.name 
                    # current_run.font.italic
                    # etc...
    
    # Replace the entire doc's font type (name) and size, save it
    if change_font: 
        print('Changing font to {} and size {}...'.format(replace_font, replace_size))
                
        for paragraph in doc.paragraphs:
            
            for r in paragraph.runs:
                r.font.name = replace_font
                r.font.size = Pt(replace_size)
                ##r.font.color.rgb = RGBColor(0, 0, 0) # set all to black
    # save as new book
    doc.save(out_book_filename)
    print('New book saved as {}'.format(out_book_filename))
    print('It took {} minutes'.format( (time.time()-t0)/60 )) # report run time
    

